import random
import zipfile
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Lists for randomizing resume content
first_names = ["Alex", "Samantha", "Michael", "Priya", "James", "Anita", "Rahul", "Emily", "David", "Neha"]
last_names = ["Smith", "Kumar", "Johnson", "Patel", "Brown", "Sharma", "Lee", "Gupta", "Wilson", "Singh"]
companies = ["TechCorp", "AI Innovations", "DataSys", "IntelliTech", "NeuroNet", "Quantum AI", "SmartAnalytics"]
universities = ["MIT", "Stanford", "UC Berkeley", "IIT Bombay", "Carnegie Mellon", "University of Toronto"]
skills = ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Deep Learning", "NLP", "Computer Vision", "AWS", "SQL"]
projects = [
    "Developed a BERT-based NLP model for sentiment analysis with 92% accuracy.",
    "Built a convolutional neural network for image classification using PyTorch.",
    "Implemented a recommendation system using collaborative filtering for e-commerce.",
    "Optimized a reinforcement learning model for autonomous navigation."
]

# Function to create a .docx resume
def create_docx_resume(name, email, phone, linkedin, years, skills, company1, company2, start_date, end_date, start_date2, end_date2, university, edu_start, edu_end, project1, project2, project3):
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Name
    name_p = doc.add_paragraph(name)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.runs[0]
    name_run.font.size = Pt(16)
    name_run.bold = True
    
    # Contact info
    contact_p = doc.add_paragraph(f"{email} | {phone} | LinkedIn: {linkedin}")
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        f"Machine Learning Engineer with {years} years of experience in designing and deploying scalable AI models. "
        f"Proficient in {skills[0]}, {skills[1]}, and {skills[2]}. Passionate about solving complex problems using data-driven approaches."
    ).style = 'Normal'
    
    # Skills
    doc.add_heading("Skills", level=1)
    skills_p = doc.add_paragraph()
    for skill in skills:
        skills_p.add_run(f"• {skill}\n").font.size = Pt(10)
    
    # Experience
    doc.add_heading("Experience", level=1)
    
    exp1 = doc.add_paragraph(f"{company1}, Machine Learning Engineer")
    exp1.add_run(f"\n{start_date} - {end_date}").italic = True
    exp1.runs[0].bold = True
    doc.add_paragraph(f"• {project1}")
    doc.add_paragraph(f"• {project2}")
    doc.add_paragraph("• Collaborated with cross-functional teams to integrate models into production.")
    
    exp2 = doc.add_paragraph(f"{company2}, Data Scientist")
    exp2.add_run(f"\n{start_date2} - {end_date2}").italic = True
    exp2.runs[0].bold = True
    doc.add_paragraph(f"• {project3}")
    doc.add_paragraph("• Developed data pipelines to preprocess large datasets for model training.")
    
    # Education
    doc.add_heading("Education", level=1)
    edu = doc.add_paragraph(f"{university}, M.S. in Computer Science")
    edu.add_run(f"\n{edu_start} - {edu_end}").italic = True
    edu.runs[0].bold = True
    
    # Projects
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(f"• {project1}")
    doc.add_paragraph(f"• {project3}")
    
    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("• Certified TensorFlow Developer")
    doc.add_paragraph("• AWS Certified Machine Learning – Specialty")
    
    return doc

# Function to generate a single resume's data
def generate_resume_data(index):
    name = f"{random.choice(first_names)} {random.choice(last_names)}"
    email = f"{name.lower().replace(' ', '.')}@example.com"
    phone = f"({random.randint(200, 999)}) {random.randint(100, 999)}-{random.randint(1000, 9999)}"
    linkedin = f"{name.lower().replace(' ', '-')}{random.randint(100, 999)}"
    years = random.randint(3, 10)
    selected_skills = random.sample(skills, 5)
    company1 = random.choice(companies)
    company2 = random.choice([c for c in companies if c != company1])
    university = random.choice(universities)
    project_list = random.sample(projects, 3)
    
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "years": years,
        "skills": selected_skills,
        "company1": company1,
        "company2": company2,
        "start_date": f"Jan {random.randint(2018, 2022)}",
        "end_date": "Present" if random.random() > 0.5 else f"Dec {random.randint(2022, 2025)}",
        "start_date2": f"Jun {random.randint(2015, 2018)}",
        "end_date2": f"Dec {random.randint(2017, 2021)}",
        "university": university,
        "edu_start": random.randint(2012, 2018),
        "edu_end": random.randint(2014, 2020),
        "project1": project_list[0],
        "project2": project_list[1],
        "project3": project_list[2]
    }

# Create a directory for resumes
output_dir = "resumes"
os.makedirs(output_dir, exist_ok=True)

# Generate 30 resumes and save them as .docx files
resume_files = []
for i in range(30):
    resume_data = generate_resume_data(i)
    doc = create_docx_resume(**resume_data)
    filename = f"{output_dir}/resume_{resume_data['name'].replace(' ', '_')}.docx"
    doc.save(filename)
    resume_files.append(filename)

# Create a zip file containing all resumes
zip_filename = "ml_engineer_resumes.zip"
with zipfile.ZipFile(zip_filename, "w", zipfile.ZIP_DEFLATED) as zipf:
    for resume_file in resume_files:
        zipf.write(resume_file)

# Clean up individual resume files (optional)
for resume_file in resume_files:
    os.remove(resume_file)
os.rmdir(output_dir)

print(f"Created {zip_filename} containing 30 .docx resumes.")